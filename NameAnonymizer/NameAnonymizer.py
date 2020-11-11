import nltk
import os
import docx
#import stanza
from itertools import count
from collections import defaultdict
from stanza.server import CoreNLPClient

os.environ["CORENLP_HOME"]="C:/Users/roma0/stanford-corenlp-4.1.0" # Point to the CoreNLP location


def preprocess(text):
    outputFile = "outputFile" + fileExtension # Define a output file with extension

    ne_person = []
    ne_person_test = []
    ne_person_list = []
    ne_name = ''
    id = 0

    # Make Coreference chain which I can use to compare elements in text
    with CoreNLPClient(
            annotators=['tokenize','ssplit','pos','lemma','ner', 'parse', 'depparse','coref'],
            timeout=30000,
            memory='16G') as client:
        ann = client.annotate(text)

    mychains = list() # Hold coreference chains
    chains = ann.corefChain
    for chain in chains:
        mychain = list()
        # Loop through every mention of this chain
        for mention in chain.mention:
            # Get the sentence in which this mention is located, and get the words which are part of this mention
            # (we can have more than one word, for example, a mention can be a pronoun like "he", but also a compound noun like "His wife Michelle")
            words_list = ann.sentence[mention.sentenceIndex].token[mention.beginIndex:mention.endIndex]

            #build a string out of the words of this mention and filter out only person        
            ment_word = ' '.join([x.word for x in words_list if x.ner == 'PERSON'])

            # Filter out repeating names and blank words
            if ment_word != '':
                if ment_word not in mychain:
                    mychain.append(ment_word)
        mychains.append(mychain)

     # Assign unique id to each coreferenc chain
    mapping = defaultdict(count().__next__)
    unique_id = []
    for element in mychains:
       unique_id.append(mapping[tuple(element)])
    
    # Replace all name occurances with a unique tag with identifier
    temp_id = 0
    for chain in mychains:
        for name in chain:
            unique_name_tag = 'PERSON' + '_' + str(unique_id[temp_id])
            text = text.replace(name, unique_name_tag)
        temp_id += 1
    
    print(text)
    ne_tree = nltk.ne_chunk(nltk.pos_tag(nltk.word_tokenize(text))) # Chunk, tokenize and label text
    #ne_tree.draw()

    for subtree in ne_tree.subtrees(filter=lambda t: t.label() == 'PERSON'): # Filter every instance of the named entity with tag PERSON
        for leaf in subtree.leaves():
            ne_person.append(leaf[0]) # Append part of full name with title
            # print (ne_person)
        for part in ne_person:
            ne_name += part + ' ' # Add part of persons name to complete it
        if ne_name[:-1] not in ne_person_list:
            ne_person_list.append(ne_name[:-1]) # Append full name to person list
            
            #print (ne_name)
        ne_name = ''
        ne_person = []
        for ne_full_name in ne_person_list: # Replace all full name occurances with a tag 'person' // Change to unique ID
            text = text.replace(ne_full_name, 'PERSON')

    #print(ne_person_list) # Print names for DEBUG PURPOSES

    newPathOut = os.path.relpath('..\\dataSet\\' + outputFile , curPath) # Make new relative path to file
    if fileExtension == ".txt":
        with open(newPathOut, 'w', encoding = 'utf-8') as outputFile: # Open file
            outputFile.write(text)
            outputFile.close()
    elif fileExtension == ".docx":
        output = docx.Document()
        output.add_paragraph(text)
        output.save(newPathOut)


text = ""
fileExtension = ".docx" #.csv ?.excel?
curProjectFile = "NameAnonymizer.sln"
inputFile = "inputFile" + fileExtension
curPath = os.path.dirname(curProjectFile) # Point to the current project folder

newPathIn = os.path.relpath('..\\dataSet\\' + inputFile , curPath)
if fileExtension == ".txt":
    with open(newPathIn, 'r', encoding = 'utf-8') as inputFile:
        text = inputFile.read()
        inputFile.close()
elif fileExtension == ".docx":
    document = docx.Document(newPathIn)
    for paragraph in document.paragraphs: 
        paragraph = paragraph.text
        text = text + paragraph

preprocess(text)


